"""Tests for DOCX full extraction."""

import os
import tempfile
from runeextract import extract


def test_docx_extraction():
    """Test basic DOCX document text extraction."""
    from docx import Document as DocxDocument
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = tmp.name
    try:
        doc = DocxDocument()
        doc.add_paragraph("Hello from python-docx.")
        doc.add_heading("Section 1", level=1)
        doc.add_paragraph("Some body text.")
        doc.save(path)

        result = extract(path)
        assert result.source_type == "docx"
        assert "Hello from python-docx" in result.text
        assert "Section 1" in result.text
    finally:
        if os.path.exists(path):
            os.unlink(path)


def test_docx_with_table():
    """Test DOCX with table extraction."""
    from docx import Document as DocxDocument
    from docx.table import Table as _  # noqa: F401
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = tmp.name
    try:
        doc = DocxDocument()
        doc.add_paragraph("Table test")
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(0, 2).text = "C"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        table.cell(1, 2).text = "3"
        doc.save(path)

        result = extract(path)
        assert len(result.tables) > 0
        if result.tables:
            assert len(result.tables[0].rows) > 0
    finally:
        if os.path.exists(path):
            os.unlink(path)


def test_docx_with_metadata():
    """Test DOCX metadata extraction."""
    from docx import Document as DocxDocument
    from docx.shared import Inches
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = tmp.name
    try:
        doc = DocxDocument()
        doc.add_paragraph("Metadata test.")
        doc.core_properties.author = "Test Author"
        doc.core_properties.title = "Test Title"
        doc.core_properties.subject = "Test Subject"
        doc.save(path)

        result = extract(path, metadata=True)
        assert result.metadata.get("title") == "Test Title"
        assert result.metadata.get("author") == "Test Author"
    finally:
        if os.path.exists(path):
            os.unlink(path)
